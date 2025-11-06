import streamlit as st
from docx import Document
from docx.shared import RGBColor
import pandas as pd
import io
import re

# Şifre
CORRECT_PASSWORD = "OxdXmX2vxM"

# Sayfa ayarları
st.set_page_config(page_title="Rapor Çeviri Sistemi", page_icon="🔬", layout="wide")

# Şifre kontrolü
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔐 Rapor Çeviri Sistemi")
    st.markdown("### Lütfen şifrenizi girin")
    
    password = st.text_input("Şifre:", type="password", key="password_input")
    
    if st.button("Giriş Yap", type="primary"):
        if password == CORRECT_PASSWORD:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Hatalı şifre! Lütfen tekrar deneyin.")
    
    st.stop()

# Excel dosyasını yükle
@st.cache_data
def load_translation_dict():
    try:
        df = pd.read_excel('Premium food&drink list_179 (1).xlsx')
        translation_dict = {}
        
        for _, row in df.iterrows():
            english = str(row.iloc[0]).strip()
            turkish = str(row.iloc[1]).strip()
            
            if english and turkish and english != 'nan' and turkish != 'nan':
                translation_dict[english.lower()] = turkish
        
        return translation_dict
    except Exception as e:
        st.error(f"Çeviri dosyası yüklenemedi: {str(e)}")
        return {}

# Çeviri fonksiyonu - Çok kelimeli ifadeleri önceliklendir
def translate_text(text, translation_dict):
    if not text or pd.isna(text):
        return text
    
    text_str = str(text).strip()
    text_lower = text_str.lower()
    
    # Önce tam eşleşme ara
    if text_lower in translation_dict:
        return translation_dict[text_lower]
    
    # Çok kelimeli ifadeleri bul ve çevir (uzundan kısaya sırala)
    sorted_keys = sorted(translation_dict.keys(), key=len, reverse=True)
    
    result = text_str
    replacements = []
    
    for key in sorted_keys:
        if len(key.split()) > 1:  # Sadece çok kelimeli ifadeler
            pattern = re.compile(re.escape(key), re.IGNORECASE)
            matches = list(pattern.finditer(result.lower()))
            
            for match in matches:
                start, end = match.span()
                replacements.append((start, end, translation_dict[key]))
    
    # Çakışmaları önlemek için sıralama
    replacements.sort(key=lambda x: x[0], reverse=True)
    
    for start, end, replacement in replacements:
        result = result[:start] + replacement + result[end:]
    
    # Tek kelimeli çeviriler
    words = result.split()
    translated_words = []
    
    for word in words:
        word_clean = word.strip('()/-,.')
        word_lower = word_clean.lower()
        
        if word_lower in translation_dict:
            prefix = word[:len(word) - len(word.lstrip('()/-,.'))]
            suffix = word[len(word.rstrip('()/-,.')):]
            translated_words.append(prefix + translation_dict[word_lower] + suffix)
        else:
            translated_words.append(word)
    
    return ' '.join(translated_words)

# DOCX çeviri fonksiyonu - FORMATLAR KORUNUYOR
def translate_docx(input_file, translation_dict):
    doc = Document(input_file)
    
    # Paragrafları çevir
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                # Orijinal formatı kaydet
                original_font = run.font
                original_bold = run.bold
                original_italic = run.italic
                original_underline = run.underline
                original_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                original_highlight = run.font.highlight_color
                original_size = run.font.size
                
                # Metni çevir
                run.text = translate_text(run.text, translation_dict)
                
                # Formatı geri yükle
                run.bold = original_bold
                run.italic = original_italic
                run.underline = original_underline
                if original_color:
                    run.font.color.rgb = original_color
                if original_highlight:
                    run.font.highlight_color = original_highlight
                if original_size:
                    run.font.size = original_size
    
    # Tabloları çevir - ARKA PLAN RENKLERİNİ KORU
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Hücre arka plan rengini kaydet
                cell_shading = cell._element.xpath('.//w:shd')
                original_fill = None
                if cell_shading:
                    original_fill = cell_shading[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            # Orijinal formatı kaydet
                            original_bold = run.bold
                            original_italic = run.italic
                            original_underline = run.underline
                            original_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                            original_highlight = run.font.highlight_color
                            original_size = run.font.size
                            
                            # Metni çevir
                            run.text = translate_text(run.text, translation_dict)
                            
                            # Formatı geri yükle
                            run.bold = original_bold
                            run.italic = original_italic
                            run.underline = original_underline
                            if original_color:
                                run.font.color.rgb = original_color
                            if original_highlight:
                                run.font.highlight_color = original_highlight
                            if original_size:
                                run.font.size = original_size
                
                # Hücre arka plan rengini geri yükle
                if original_fill and cell_shading:
                    cell_shading[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', original_fill)
    
    # Belleğe kaydet
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Ana uygulama
st.title("🔬 Premium Food Intolerance Test - Rapor Çeviri Sistemi")
st.markdown("### İngilizce raporları Türkçe'ye çevirin")

# Çeviri sözlüğünü yükle
translation_dict = load_translation_dict()

if translation_dict:
    st.success(f"✅ {len(translation_dict)} çeviri yüklendi!")
    
    # Dosya yükleme
    uploaded_file = st.file_uploader(
        "DOCX dosyasını yükleyin",
        type=['docx'],
        help="Sadece .docx formatındaki dosyalar desteklenmektedir"
    )
    
    if uploaded_file:
        st.info(f"📄 Dosya: **{uploaded_file.name}**")
        
        if st.button("🚀 Çeviriyi Başlat", type="primary"):
            with st.spinner("Çeviriliyor... Lütfen bekleyin..."):
                try:
                    # Çeviri yap
                    translated_file = translate_docx(uploaded_file, translation_dict)
                    
                    # İndirme butonu
                    st.success("✅ Çeviri tamamlandı!")
                    
                    output_filename = uploaded_file.name.replace('.docx', '_TR.docx')
                    
                    st.download_button(
                        label="📥 Çevrilmiş Dosyayı İndir",
                        data=translated_file,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"❌ Hata oluştu: {str(e)}")
else:
    st.error("❌ Çeviri dosyası yüklenemedi!")

# Çıkış butonu
st.sidebar.markdown("---")
if st.sidebar.button("🚪 Çıkış Yap"):
    st.session_state.authenticated = False
    st.rerun()
